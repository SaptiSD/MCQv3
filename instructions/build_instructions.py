"""Generate questions/instructions docx for the MCQ platform."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "MCQv3 Instructions.docx"

GREEN = RGBColor(0x17, 0x6B, 0x52)
MUTED = RGBColor(0x53, 0x63, 0x5D)
INK = RGBColor(0x15, 0x23, 0x1F)


def heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GREEN
    return h


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def steps(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = MUTED


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for style_name in ("Normal", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = INK

    title = doc.add_heading("MCQ Assessment Platform — User Instructions", level=0)
    for run in title.runs:
        run.font.color.rgb = INK
    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run("How to sign in and use the platform as a teacher, a student, or an administrator.")
    tagline_run.font.color.rgb = MUTED
    tagline_run.italic = True
    doc.add_page_break()

    heading(doc, "1.  Signing in", 1)
    para(doc, "Open the platform in a browser. On the login screen, enter your Email or full name and your password, then click Login. Role determines what you see:")
    bullets(doc, [
        "Teachers and students sign in with their email address or their first and last name.",
        "Administrators sign in with their admin email address.",
        "Use the Sign out button at the top to switch accounts.",
    ])
    callout(doc, "Demo accounts: teacher — teacher@mcq.local / teacher · student — student@mcq.local / student · admin — admin@mcq.local / admin")

    heading(doc, "2.  For Teachers", 1)
    para(doc, "The teacher workspace has five pages in the top navigation: Dashboard, Create quiz, Students, Analytics, and Student view.")

    heading(doc, "2.1  Dashboard", 2)
    para(doc, "Your overview page. It shows key figures (total quizzes, published, students in roster, completed attempts, average score, pass rate, assigned students, completion rate) and lists every quiz you own.")
    bullets(doc, [
        "Create new — start building a quiz.",
        "Search quizzes — filter your list by title.",
        "Manage — open a quiz to edit it, change its settings, or review attempts.",
        "Delete — remove an assessment (a confirmation dialog asks first).",
    ])

    heading(doc, "2.2  Create quiz", 2)
    steps(doc, [
        "Open Create quiz and click Publish quiz at the top or the bottom of the page.",
        "Questions — add at least one question, two ways:",
    ])
    bullets(doc, [
        "Create manually — pick a question type (Multiple choice, Multiple choice - select all that apply, True / False, Fill in the blank, Short answer), enter the question text and options, and mark the correct answer. Use Add another question / Remove last question to size the quiz.",
        "Upload question bank — upload a .txt or .docx bank and click Read question bank. It must use numbered questions, options, and an Answer Key so the platform can parse it.",
    ])
    steps(doc, [
        "Settings — give the quiz a title, set time allowed (minutes), passing score (%), and choose options:",
    ])
    bullets(doc, [
        "Allow retakes — students may take the quiz again after finishing.",
        "Show class average to students — lets students see the class average on results.",
        "Randomize question order / Randomize answer order — shuffle questions and optional choices for each student (on by default).",
        "Opening / closing — set an available window (or leave it open).",
    ])
    steps(doc, [
        "Audience — choose which students take it: assign to specific students, assign to teams, or assign nobody (the quiz is then open to all students).",
        "Click Publish quiz. Problems (missing title, missing questions, invalid closing time) are reported before publishing.",
    ])

    heading(doc, "2.3  Manage a quiz", 2)
    para(doc, "From the Dashboard, click Manage on a quiz to edit its questions and settings, change the assignment, and see who has submitted with their scores and results. The dashboard also shows each quiz's question count, duration, passing score, and audience.")

    heading(doc, "2.4  Students", 2)
    bullets(doc, [
        "Roster — every student in your workspace with their assignment and performance.",
        "Add a student — enter a name and email; new addresses get a student account automatically.",
        "Teams — create teams and add members, then assign quizzes to whole teams.",
        "Performance — averages, pass rates, and last activity, with an export.",
    ])

    heading(doc, "2.5  Analytics", 2)
    para(doc, "Per-quiz progress shows every assigned student with status (Not started / In progress / Completed), score, result (Passed / Failed), and last activity, plus search and an export.")

    heading(doc, "2.6  Student view", 2)
    para(doc, "Open Student view to experience the student-side exactly as your students will — see assignments and take a quiz from the student's point of view.")

    heading(doc, "3.  For Students", 1)
    steps(doc, [
        "Sign in with your email or full name.",
        "You land on the Student workspace. Every quiz assigned and open right now appears with its duration and closing date, together with your latest result.",
        "Click Start quiz to begin, Resume to continue a quiz you already opened, or Retake if your teacher allows retakes.",
        "Answer each question and use Submit at the end. If the time limit or closing time runs out first, the attempt is submitted automatically.",
        "You see your score right away with Passed or Needs another try — and the class average when your teacher shows it.",
    ])
    callout(doc, "If you close the page mid-quiz, the next time you return your unfinished attempt is still open and you can Resume — the clock keeps running.")

    heading(doc, "4.  For Administrators", 1)
    para(doc, "Sign in with an administrator email to open the Admin console, which manages every account on the platform in three sections: Teachers, Students, and Administrators.")

    heading(doc, "4.1  Teachers and Students", 2)
    steps(doc, [
        "Add — enter Name, Email, and Password, then click Add teacher or Add student.",
        "Edit — select an account, change its name, email, or new password, and Save changes.",
        "Remove — select the account and remove it. Accounts that own quizzes, teams, or attempt history cannot be removed.",
    ])

    heading(doc, "4.2  Administrators", 2)
    steps(doc, [
        "Add — enter Email, Password, and Name, then click Add administrator.",
        "Edit — select an administrator, update email, name, or new password, and Save changes.",
        "Remove — select an administrator and remove them. You cannot remove your own account, and at least one administrator must always remain.",
    ])

    callout(doc, "Every account change (add, edit, or remove) applies immediately across the platform.")

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()