"""Teacher-facing assessment management screens."""

from __future__ import annotations

import io
import json
from datetime import datetime, time, timedelta, timezone

import pandas as pd
import streamlit as st
from docx import Document

from db import connect
from ingestion import extract_upload, parse_bank
from repository import (add_student_to_roster, assigned_student_ids, create_quiz,
                        delete_quiz, move_question, questions_for_quiz, quiz_for_teacher, quizzes_for_teacher,
                        save_question_bank, set_quiz_assignments, students,
                        student_analytics, student_detail_analytics, student_progress_for_quiz,
                        set_team_members, student_ids_for_teams, team_student_ids, teams_for_student, teams_for_teacher,
                        teacher_analytics, update_quiz_settings)


SELECT_ALL_TYPE = "Multiple choice - select all that apply"
QUESTION_TYPES = ["Multiple choice", SELECT_ALL_TYPE, "True / False", "Fill in the blank", "Short answer"]


@st.dialog("Delete assessment?")
def delete_quiz_dialog(user, quiz_id: int, title: str) -> None:
    st.write(f"Delete **{title}** and all its questions, assignments, and results?")
    confirm, cancel = st.columns(2)
    if confirm.button("Yes, delete", key=f"confirm-delete-{quiz_id}", type="primary", width="stretch"):
        delete_quiz(quiz_id, user["id"])
        st.session_state.pop("manage_quiz", None)
        st.rerun()
    if cancel.button("Cancel", key=f"cancel-delete-{quiz_id}", width="stretch"):
        st.rerun()


def dashboard(user) -> None:
    quizzes = quizzes_for_teacher(user["id"])
    roster = students(user["id"])
    analytics = teacher_analytics(user["id"])
    st.markdown('<div class="eyebrow">Teacher workspace</div><h1>Your assessments</h1>', unsafe_allow_html=True)
    st.caption("Create, assign, review, and understand the assessments you own.")
    metrics = [(analytics["quizzes"], "Total quizzes"), (analytics["active_quizzes"], "Published"), (len(roster), "Students in roster"), (analytics["completed"], "Completed attempts"), (f"{analytics['average_score'] or 0:.1f}%", "Average score"), (f"{(analytics['pass_rate'] or 0) * 100:.0f}%", "Pass rate"), (analytics["assigned_students"], "Assigned students"), (f"{(analytics['completed'] / analytics['attempts'] * 100) if analytics['attempts'] else 0:.0f}%", "Completion rate")]
    for row in range(0, len(metrics), 4):
        for column, (value, label) in zip(st.columns(4), metrics[row:row + 4]):
            with column: st.markdown(f'<div class="metric"><strong>{value}</strong><small>{label}</small></div>', unsafe_allow_html=True)
    st.divider()
    create_col, search_col = st.columns([1, 3], vertical_alignment="center")
    if create_col.button("＋  Create new", type="primary", width="stretch"):
        st.session_state.page_override = "Create quiz"
        st.rerun()
    quiz_search = search_col.text_input("Search quizzes", placeholder="Search by title", label_visibility="collapsed", key="dashboard-quiz-search")
    visible_quizzes = [quiz for quiz in quizzes if not quiz_search.strip() or quiz_search.lower() in quiz["title"].lower()]
    if not visible_quizzes:
        st.info("No assessments match your search.")
        return
    for quiz in visible_quizzes:
        with st.container(border=True):
            details, action = st.columns([4, 1])
            with details:
                st.subheader(quiz["title"])
                assigned = len(assigned_student_ids(quiz["id"]))
                audience = f"{assigned} assigned students" if assigned else "All students"
                st.caption(f"{len(questions_for_quiz(quiz['id']))} questions  ·  {quiz['duration_minutes']} minutes  ·  pass at {quiz['passing_score']}%  ·  {audience}")
            with action:
                if st.button("Manage", key=f"manage-{quiz['id']}", width="stretch"):
                    st.session_state.manage_quiz = quiz["id"]; st.rerun()
                if st.button("Delete", key=f"delete-{quiz['id']}", width="stretch"):
                    delete_quiz_dialog(user, quiz["id"], quiz["title"])
        if st.session_state.get("manage_quiz") == quiz["id"]:
            with st.container(border=True):
                manage_quiz(user, quiz["id"])


def analytics_page(user) -> None:
    analytics = teacher_analytics(user["id"])
    students_data = student_analytics(user["id"])
    st.markdown('<div class="eyebrow">Teacher workspace</div><h1>Performance overview</h1>', unsafe_allow_html=True)
    st.caption("A quick read on assessment health, student outcomes, and completion.")
    metrics = [(analytics["attempts"], "Total attempts"), (analytics["completed"], "Completed"), (f"{analytics['average_score'] or 0:.1f}%", "Average score"), (f"{(analytics['pass_rate'] or 0) * 100:.0f}%", "Pass rate"), (f"{(analytics['completed'] / analytics['attempts'] * 100) if analytics['attempts'] else 0:.0f}%", "Completion rate"), (len(students_data), "Students tracked")]
    for column, (value, label) in zip(st.columns(6), metrics):
        with column: st.markdown(f'<div class="metric"><strong>{value}</strong><small>{label}</small></div>', unsafe_allow_html=True)
    st.divider()
    st.subheader("Student performance")
    if students_data:
        st.dataframe(pd.DataFrame([{"Student": row["name"], "Email": row["email"], "Assigned": row["assigned_quizzes"], "Attempts": row["attempts"], "Completed": row["completed"], "Average score": f"{row['average_score']:.1f}%" if row["average_score"] is not None else "-", "Pass rate": f"{row['pass_rate'] * 100:.0f}%" if row["pass_rate"] is not None else "-", "Last activity": row["last_activity"] or "-"} for row in students_data]), width="stretch", hide_index=True)
    else:
        st.info("Add students to your roster to start tracking performance.")
    st.divider()
    st.subheader("Exam participation")
    quizzes = quizzes_for_teacher(user["id"])
    if quizzes:
        quiz_options = {quiz["id"]: quiz["title"] for quiz in quizzes}
        selected_quiz_id = st.selectbox("Select an exam", list(quiz_options), format_func=quiz_options.get, key="analytics-exam")
        progress = student_progress_for_quiz(user["id"], selected_quiz_id)
        if progress:
            search = st.text_input("Look up a student on this test", placeholder="Search by name or email", key="analytics-student-search")
            if search.strip():
                progress = [row for row in progress if search.lower() in row["student"].lower() or search.lower() in row["email"].lower()]
            counts = {status: sum(row["status"] == status for row in progress) for status in ("Not started", "In progress", "Completed")}
            for column, (value, label) in zip(st.columns(3), [(counts["Not started"], "Not started"), (counts["In progress"], "In progress"), (counts["Completed"], "Completed")]):
                with column: st.markdown(f'<div class="metric"><strong>{value}</strong><small>{label}</small></div>', unsafe_allow_html=True)
            st.dataframe(pd.DataFrame([{"Student": row["student"], "Email": row["email"], "Status": row["status"], "Score": f"{row['score']:.1f}%" if row["score"] is not None else "-", "Result": row["result"], "Last activity": row["last_activity"]} for row in progress]), width="stretch", hide_index=True)
        else:
            st.info("No students are assigned to this exam yet.")


def _create_save_setting(key: str) -> None:
    if "new_quiz_data" not in st.session_state:
        st.session_state["new_quiz_data"] = {}
    st.session_state["new_quiz_data"][key] = st.session_state[key]


def create(user) -> None:
    form = st.session_state.setdefault("new_quiz_data", {})
    st.markdown('<div class="eyebrow">New assessment</div><h1>Shape the experience</h1>', unsafe_allow_html=True)
    st.caption("Build questions first, set the rules second, then publish when everything is ready.")
    top_publish = st.button("Publish quiz", key="new-quiz-publish-top", type="primary", width="stretch")
    section = st.session_state.get("new-quiz-section", "questions")
    settings_button, questions_button = st.columns(2)
    if settings_button.button("Quiz settings", key="new-quiz-settings", type="primary" if section == "settings" else "secondary", width="stretch"):
        st.session_state["new-quiz-section"] = "settings"
        st.rerun()
    if questions_button.button("Questions", key="new-quiz-questions", type="primary" if section == "questions" else "secondary", width="stretch"):
        st.session_state["new-quiz-section"] = "questions"
        st.rerun()

    if section == "questions":
        with st.container(border=True):
            st.subheader("Questions")
            question_mode = st.radio("Add questions", ["Create manually", "Upload question bank"], horizontal=True, key="new-quiz-mode", on_change=_create_save_setting, args=("new-quiz-mode",))
            if question_mode == "Upload question bank":
                upload = st.file_uploader("Question bank (.txt or .docx)", type=["txt", "docx"], key="new-quiz-upload")
                if upload and st.button("Read question bank", key="new-quiz-parse"):
                    try:
                        form["new-uploaded-questions"] = parse_bank(extract_upload(upload))
                    except Exception as exc:
                        form["new-uploaded-questions"] = []
                        st.error(f"Could not read this question bank: {exc}")
                uploaded_questions = form.get("new-uploaded-questions", [])
                if not isinstance(uploaded_questions, list):
                    uploaded_questions = []
                st.write(f"Questions ready: **{len(uploaded_questions)}**")
                if upload and not uploaded_questions:
                    st.warning("No valid questions found. Include numbered questions, options, and an Answer Key before publishing.")
            else:
                count_key = "new-manual-count"
                if count_key not in form:
                    form[count_key] = 1
                count = form[count_key]
                add_col, remove_col = st.columns(2)
                st.write(f"**{int(count)}** question{'s' if int(count) != 1 else ''} in this quiz")
                if add_col.button("Add another question", key="new-manual-add", width="stretch"):
                    form[count_key] = int(count) + 1
                    st.rerun()
                if remove_col.button("Remove last question", key="new-manual-remove", width="stretch", disabled=int(count) <= 1):
                    form[count_key] = int(count) - 1
                    st.rerun()
                for index in range(int(count)):
                    with st.container(border=True):
                        st.markdown(f"**Question {index + 1}**")
                        current_type = form.get(f"new-type-{index}", "Multiple choice")
                        question_type = st.selectbox("Question type", QUESTION_TYPES, index=QUESTION_TYPES.index(current_type) if current_type in QUESTION_TYPES else 0, key=f"new-type-{index}", on_change=_create_save_setting, args=(f"new-type-{index}",))
                        st.text_area("Question text", key=f"new-text-{index}", height=80, value=form.get(f"new-text-{index}", ""), on_change=_create_save_setting, args=(f"new-text-{index}",))
                        if question_type in {"Multiple choice", SELECT_ALL_TYPE}:
                            option_cols = st.columns(4)
                            for option_index, label in enumerate(("A", "B", "C", "D")):
                                with option_cols[option_index]:
                                    st.text_input(f"Option {label}", key=f"new-option-{index}-{label}", value=form.get(f"new-option-{index}-{label}", ""), on_change=_create_save_setting, args=(f"new-option-{index}-{label}",))
                            if question_type == SELECT_ALL_TYPE:
                                st.multiselect("Correct answers", ["A", "B", "C", "D"], key=f"new-correct-all-{index}", default=form.get(f"new-correct-all-{index}", []), on_change=_create_save_setting, args=(f"new-correct-all-{index}",))
                            else:
                                correct_cfg = ["A", "B", "C", "D"]
                                correct_val = form.get(f"new-correct-{index}")
                                st.selectbox("Correct answer", correct_cfg, index=(correct_cfg.index(correct_val) if correct_val in correct_cfg else 0), key=f"new-correct-{index}", on_change=_create_save_setting, args=(f"new-correct-{index}",))
                        elif question_type == "True / False":
                            tf_val = form.get(f"new-correct-{index}")
                            st.selectbox("Correct answer", ["True", "False"], index=(0 if tf_val != "False" else 1), key=f"new-correct-{index}", on_change=_create_save_setting, args=(f"new-correct-{index}",))
                        else:
                            st.text_input("Correct answer", key=f"new-correct-{index}", value=form.get(f"new-correct-{index}", ""), on_change=_create_save_setting, args=(f"new-correct-{index}",))
    else:
        with st.container(border=True):
            st.subheader("Quiz settings")
            st.text_input("Quiz title", placeholder="e.g. Foundations of Computing", key="new-title",
                          value=form.get("new-title", ""), on_change=_create_save_setting, args=("new-title",))
            first, second = st.columns(2)
            with first: st.number_input("Time allowed (minutes)", 1, 480, form.get("new-duration", 30), key="new-duration", on_change=_create_save_setting, args=("new-duration",))
            with second: st.number_input("Passing score (%)", 0, 100, form.get("new-passing", 70), key="new-passing", on_change=_create_save_setting, args=("new-passing",))
            st.checkbox("Allow retakes", form.get("new-retakes", False), key="new-retakes", on_change=_create_save_setting, args=("new-retakes",))
            st.checkbox("Show class average to students", form.get("new-average", False), key="new-average", on_change=_create_save_setting, args=("new-average",))
            st.checkbox("Randomize question order", form.get("new-randomize-questions", True), key="new-randomize-questions", on_change=_create_save_setting, args=("new-randomize-questions",))
            st.checkbox("Randomize answer order", form.get("new-randomize-answers", True), key="new-randomize-answers", on_change=_create_save_setting, args=("new-randomize-answers",))
            today = datetime.now().date()
            tomorrow = today + timedelta(days=1)
            opening_enabled = form.get("new-opening-enabled", True)
            st.checkbox("Enable opening date and time", opening_enabled, key="new-opening-enabled", on_change=_create_save_setting, args=("new-opening-enabled",))
            opening_date, opening_time = st.columns(2)
            with opening_date: st.date_input("Opens on", form.get("new-opening-day", today), key="new-opening-day", disabled=not opening_enabled, on_change=_create_save_setting, args=("new-opening-day",))
            with opening_time: st.time_input("Opening time", form.get("new-opening-clock", time(8, 0)), key="new-opening-clock", disabled=not opening_enabled, on_change=_create_save_setting, args=("new-opening-clock",))
            closing_enabled = form.get("new-closing-enabled", True)
            st.checkbox("Enable closing date and time", closing_enabled, key="new-closing-enabled", on_change=_create_save_setting, args=("new-closing-enabled",))
            closing_date, closing_time = st.columns(2)
            with closing_date: st.date_input("Closes on", form.get("new-closing-day", tomorrow), key="new-closing-day", disabled=not closing_enabled, on_change=_create_save_setting, args=("new-closing-day",))
            with closing_time: st.time_input("Closing time", form.get("new-closing-clock", time(17, 0)), key="new-closing-clock", disabled=not closing_enabled, on_change=_create_save_setting, args=("new-closing-clock",))
            st.divider(); st.subheader("Audience")
            roster = students(user["id"])
            audience_mode = st.radio("Assign to", ["Students", "Teams"], horizontal=True, key="new-audience-mode",
                                     index=0 if form.get("new-audience-mode") != "Teams" else 1, on_change=_create_save_setting, args=("new-audience-mode",))
            if audience_mode == "Students":
                st.multiselect("Assign to students", options=roster, default=form.get("new-selected", []), format_func=lambda row: f"{row['name']}  ·  {row['email']}", key="new-selected", on_change=_create_save_setting, args=("new-selected",))
            else:
                st.multiselect("Assign to teams", options=teams_for_teacher(user["id"]), default=form.get("new-team-selected", []), format_func=lambda team: team["name"], key="new-team-selected", on_change=_create_save_setting, args=("new-team-selected",))

    bottom_publish = st.button("Publish quiz", key="new-quiz-publish-bottom", type="primary", width="stretch")
    if not (top_publish or bottom_publish):
        return
    title = form.get("new-title", "").strip()
    question_mode = form.get("new-quiz-mode", "Create manually")
    if question_mode == "Upload question bank":
        questions = form.get("new-uploaded-questions", [])
        if not isinstance(questions, list):
            questions = []
    else:
        questions = []
        for index in range(int(form.get("new-manual-count", 1))):
            question_type = form.get(f"new-type-{index}", "Multiple choice")
            if question_type in {"Multiple choice", SELECT_ALL_TYPE}:
                options = [(label, form.get(f"new-option-{index}-{label}", "").strip()) for label in ("A", "B", "C", "D")]
                options = [(label, value) for label, value in options if value]
                correct = form.get(f"new-correct-all-{index}", []) if question_type == SELECT_ALL_TYPE else form.get(f"new-correct-{index}", "A")
            elif question_type == "True / False":
                options = [("A", "True"), ("B", "False")]
                correct = "A" if form.get(f"new-correct-{index}") == "True" else "B"
            else:
                answer = form.get(f"new-correct-{index}", "").strip()
                options, correct = ([("A", answer)] if answer else []), "A"
            questions.append({"question_text": form.get(f"new-text-{index}", "").strip(), "options": options, "correct_label": correct, "question_type": question_type})
    errors = []
    if not title:
        errors.append("Give the quiz a title first.")
    if not questions:
        errors.append("Add at least one question before publishing.")
    for index, question in enumerate(questions, 1):
        minimum = 1 if question.get("question_type") in {"Fill in the blank", "Short answer"} else 2
        if not question["question_text"]:
            errors.append(f"Question {index} needs text.")
        if len(question["options"]) < minimum:
            errors.append(f"Question {index} needs an answer." if minimum == 1 else f"Question {index} needs at least two options.")
        correct_labels = question["correct_label"] if question.get("question_type") == SELECT_ALL_TYPE else [question["correct_label"]]
        if question.get("question_type") == SELECT_ALL_TYPE and not correct_labels:
            errors.append(f"Question {index} needs at least one correct answer.")
        if not set(correct_labels).issubset({label for label, _ in question["options"]}):
            errors.append(f"Question {index} needs its selected correct option filled in.")
    if errors:
        st.error(" ".join(errors))
        return
    opening = datetime.combine(form.get("new-opening-day", datetime.now().date()), form.get("new-opening-clock", time(8, 0)), tzinfo=timezone.utc)
    closing = datetime.combine(form.get("new-closing-day", (datetime.now().date() + timedelta(days=1))), form.get("new-closing-clock", time(17, 0)), tzinfo=timezone.utc)
    opening_enabled = form.get("new-opening-enabled", True)
    closing_enabled = form.get("new-closing-enabled", True)
    now = datetime.now(timezone.utc)
    if opening_enabled and closing_enabled and closing <= opening:
        st.error("Closing time must be after opening time.")
        return
    if closing_enabled and closing <= now:
        st.error("Closing time is in the past; students won't be able to take this quiz. Set a closing time in the future.")
        return
    if opening_enabled and opening > now:
        st.info(f"The quiz opens on {opening.strftime('%b %d, %I:%M %p')} and won't be visible to students until then.")
    assigned_students = {row["id"] for row in form.get("new-selected", [])}
    assigned_students.update(student_ids_for_teams(user["id"], [team["id"] for team in form.get("new-team-selected", [])]))
    quiz_id = create_quiz(user["id"], title, form.get("new-duration", 30), form.get("new-passing", 70), form.get("new-retakes", False), form.get("new-average", False), opening.isoformat(), closing.isoformat(), list(assigned_students), opening_enabled, closing_enabled, form.get("new-randomize-questions", True), form.get("new-randomize-answers", True))
    save_question_bank(quiz_id, questions)
    st.session_state.manage_quiz = quiz_id
    st.session_state[f"quiz-section-{quiz_id}"] = "questions"
    st.session_state.pop("new-quiz-section", None)
    st.session_state.pop("new_quiz_data", None)
    st.rerun()


def manage_quiz(user, quiz_id: int) -> None:
    quiz = quiz_for_teacher(quiz_id, user["id"])
    if not quiz: return
    st.divider(); st.markdown(f"### Manage: {quiz['title']}")
    questions_button, settings_button = st.columns(2)
    section = st.session_state.get(f"quiz-section-{quiz_id}", "questions")
    if settings_button.button("Quiz settings", key=f"settings-section-{quiz_id}", type="primary" if section == "settings" else "secondary", width="stretch"):
        st.session_state[f"quiz-section-{quiz_id}"] = "settings"
        st.rerun()
    if questions_button.button("Questions", key=f"questions-section-{quiz_id}", type="primary" if section == "questions" else "secondary", width="stretch"):
        st.session_state[f"quiz-section-{quiz_id}"] = "questions"
        st.rerun()
    if section == "settings":
        settings_editor(quiz)
    else:
        question_bank(quiz)
    with st.expander("Assign by students or teams · View results"):
        assignment_editor(quiz)
        results(quiz)
    if st.button("Close manager", key=f"close-{quiz_id}"): st.session_state.pop("manage_quiz", None); st.rerun()


def question_bank(quiz) -> None:
    questions = questions_for_quiz(quiz["id"])
    if questions:
        with st.expander("Reorder questions"):
            question_options = {question["id"]: f"{index}. {question['question_text']}" for index, question in enumerate(questions, 1)}
            selected_id = st.selectbox("Question", list(question_options), format_func=question_options.get, key=f"reorder-question-{quiz['id']}")
            selected_index = next(index for index, question in enumerate(questions) if question["id"] == selected_id)
            move_up, move_down = st.columns(2)
            if move_up.button("Move up", key=f"move-up-{quiz['id']}", disabled=selected_index == 0, width="stretch"):
                move_question(quiz["id"], selected_id, -1)
                st.rerun()
            if move_down.button("Move down", key=f"move-down-{quiz['id']}", disabled=selected_index == len(questions) - 1, width="stretch"):
                move_question(quiz["id"], selected_id, 1)
                st.rerun()
    mode = st.radio("How would you like to add questions?", ["Create manually", "Upload question bank"], horizontal=True, key=f"question-mode-{quiz['id']}")
    if mode == "Create manually":
        manual_question_editor(quiz)
        return
    upload = st.file_uploader("Upload question bank (.txt or .docx)", type=["txt", "docx"], key=f"upload-{quiz['id']}")
    if upload and st.button("Parse question bank", type="primary", key=f"parse-{quiz['id']}"):
        try:
            st.session_state[f"draft-{quiz['id']}"] = parse_bank(extract_upload(upload))
        except Exception as exc:
            st.session_state[f"draft-{quiz['id']}"] = []
            st.error(f"Could not read this question bank: {exc}")
    draft = st.session_state.get(f"draft-{quiz['id']}")
    if draft is not None:
        if not draft:
            st.warning("No valid questions found. Include numbered questions, options, and an Answer Key before publishing.")
            return
        st.write("Review extracted questions before publishing")
        table = pd.DataFrame([{"Question": q["question_text"], "Options": " | ".join(f"{a}) {b}" for a, b in q["options"]), "Correct": q["correct_label"]} for q in draft])
        edited = st.data_editor(table, num_rows="dynamic", width="stretch", key=f"editor-{quiz['id']}")
        if st.button("Save question bank and publish", type="primary", key=f"save-{quiz['id']}"):
            questions = []
            for _, row in edited.iterrows():
                options = [(chr(65 + i), part.split(")", 1)[-1].strip()) for i, part in enumerate(str(row["Options"]).split("|"))]
                questions.append({"question_text": str(row["Question"]), "options": options, "correct_label": str(row["Correct"]).strip().upper()})
            save_question_bank(quiz["id"], questions); st.session_state.pop(f"draft-{quiz['id']}", None); st.success("Question bank published."); st.rerun()
        return
    st.caption(f"Question bank: {len(questions)} questions")
    for index, question in enumerate(questions, 1):
        options = json.loads(question["options_json"])
        st.write(f"**{index}. {question['question_text']}**")
        st.caption("  ·  ".join(f"{label}) {text}" for label, text in options))
    if questions:
        frame = pd.DataFrame([{"Question": q["question_text"], "Correct": q["correct_label"], **dict(json.loads(q["options_json"]))} for q in questions])
        st.download_button("Download question bank CSV", frame.to_csv(index=False), "question-bank.csv", "text/csv", key=f"bank-csv-{quiz['id']}")
        document = Document(); document.add_heading(quiz["title"], 0)
        for index, q in enumerate(questions, 1):
            document.add_paragraph(f"{index}. {q['question_text']}")
            for label, text in json.loads(q["options_json"]): document.add_paragraph(f"{label}) {text}", style="List Bullet")
        output = io.BytesIO(); document.save(output)
        st.download_button("Download printable DOCX", output.getvalue(), "quiz.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"docx-{quiz['id']}")


def manual_question_editor(quiz) -> None:
    st.caption("Create the test directly. Each question needs text, at least two options, and one correct answer.")
    count_key = f"manual-count-{quiz['id']}"
    if count_key not in st.session_state:
        existing_questions = list(questions_for_quiz(quiz["id"]))
        st.session_state[count_key] = max(1, len(existing_questions))
        for index, question in enumerate(existing_questions):
            question_type = question["question_type"]
            st.session_state[f"manual-type-{quiz['id']}-{index}"] = question_type
            st.session_state[f"manual-text-{quiz['id']}-{index}"] = question["question_text"]
            for label, value in json.loads(question["options_json"]):
                st.session_state[f"manual-option-{quiz['id']}-{index}-{label}"] = value
            if question_type == SELECT_ALL_TYPE:
                st.session_state[f"manual-correct-all-{quiz['id']}-{index}"] = json.loads(question["correct_label"])
            elif question_type in {"Fill in the blank", "Short answer"}:
                options = json.loads(question["options_json"])
                st.session_state[f"manual-correct-{quiz['id']}-{index}"] = options[0][1] if options else ""
            else:
                st.session_state[f"manual-correct-{quiz['id']}-{index}"] = question["correct_label"]
    count = st.number_input("Number of questions", min_value=1, max_value=200, value=st.session_state[count_key])
    add_col, remove_col = st.columns(2)
    if add_col.button("Add another question", key=f"manual-add-{quiz['id']}", width="stretch"):
        st.session_state[count_key] = int(count) + 1
        st.rerun()
    if remove_col.button("Remove last question", key=f"manual-remove-{quiz['id']}", width="stretch", disabled=int(count) <= 1):
        st.session_state[count_key] = int(count) - 1
        st.rerun()
    questions = []
    for index in range(int(count)):
        with st.container(border=True):
            st.markdown(f"**Question {index + 1}**")
            question_type = st.selectbox("Question type", QUESTION_TYPES, key=f"manual-type-{quiz['id']}-{index}")
            text = st.text_area("Question text", key=f"manual-text-{quiz['id']}-{index}", height=80)
            options = []
            if question_type == "True / False":
                options = [("A", "True"), ("B", "False")]
            elif question_type in {"Multiple choice", SELECT_ALL_TYPE}:
                columns = st.columns(4)
                for option_index, label in enumerate(("A", "B", "C", "D")):
                    with columns[option_index]:
                        value = st.text_input(f"Option {label}", key=f"manual-option-{quiz['id']}-{index}-{label}")
                        if value.strip():
                            options.append((label, value.strip()))
            if question_type == "True / False":
                correct = st.selectbox("Correct answer", ["A", "B"], format_func=lambda value: "True" if value == "A" else "False", key=f"manual-correct-{quiz['id']}-{index}")
            elif question_type == SELECT_ALL_TYPE:
                correct = st.multiselect("Correct answers", ["A", "B", "C", "D"], key=f"manual-correct-all-{quiz['id']}-{index}")
            else:
                correct = st.text_input("Correct answer", value="A" if question_type == "Multiple choice" else "", key=f"manual-correct-{quiz['id']}-{index}")
            if question_type in {"Fill in the blank", "Short answer"}:
                options = [("A", correct.strip())] if correct.strip() else []
            correct_label = "A" if question_type in {"Fill in the blank", "Short answer"} else (correct if question_type == SELECT_ALL_TYPE else correct.strip().upper())
            questions.append({"question_text": text.strip(), "options": options, "correct_label": correct_label, "question_type": question_type})
    if st.button("Save manually created test", type="primary", key=f"manual-save-{quiz['id']}", width="stretch"):
        errors = []
        for index, question in enumerate(questions, 1):
            if not question["question_text"]:
                errors.append(f"Question {index} needs text.")
            minimum_options = 1 if question["question_type"] in {"Fill in the blank", "Short answer"} else 2
            if len(question["options"]) < minimum_options:
                errors.append(f"Question {index} needs an answer." if minimum_options == 1 else f"Question {index} needs at least two options.")
            correct_labels = question["correct_label"] if question["question_type"] == SELECT_ALL_TYPE else [question["correct_label"]]
            if question["question_type"] == SELECT_ALL_TYPE and not correct_labels:
                errors.append(f"Question {index} needs at least one correct answer.")
            if not set(correct_labels).issubset({label for label, _ in question["options"]}):
                errors.append(f"Question {index} needs its selected correct option filled in.")
        if errors:
            st.error(" ".join(errors))
        else:
            save_question_bank(quiz["id"], questions)
            st.success("Test saved and published.")
            st.rerun()


def settings_editor(quiz) -> None:
    with connect() as db:
        has_attempts = db.execute("SELECT 1 FROM attempts WHERE quiz_id = ? LIMIT 1", (quiz["id"],)).fetchone() is not None
    if has_attempts:
        st.info("Settings are read-only after a student starts this assessment.")
    opening = datetime.fromisoformat(quiz["opening_time"])
    closing = datetime.fromisoformat(quiz["closing_time"])
    with st.form(f"settings-{quiz['id']}"):
        first, second, third = st.columns(3)
        with first: duration = st.number_input("Time allowed (minutes)", 1, 480, quiz["duration_minutes"], disabled=has_attempts)
        with second: passing = st.number_input("Passing score (%)", 0, 100, quiz["passing_score"], disabled=has_attempts)
        opening_enabled = st.checkbox("Enable opening date and time", value=bool(quiz["opening_enabled"]), disabled=has_attempts)
        with st.container(border=True):
            opening_date, opening_time = st.columns(2)
            with opening_date: opening_day = st.date_input("Opens on", opening.date(), disabled=has_attempts or not opening_enabled)
            with opening_time: opening_clock = st.time_input("Opening time", opening.timetz().replace(tzinfo=None), disabled=has_attempts or not opening_enabled)
        closing_enabled = st.checkbox("Enable closing date and time", value=bool(quiz["closing_enabled"]), disabled=has_attempts)
        with st.container(border=True):
            closing_date, closing_time = st.columns(2)
            with closing_date: closing_day = st.date_input("Closes on", closing.date(), disabled=has_attempts or not closing_enabled)
            with closing_time: closing_clock = st.time_input("Closing time", closing.timetz().replace(tzinfo=None), disabled=has_attempts or not closing_enabled)
        allow_retake = st.checkbox("Allow retakes", bool(quiz["allow_retake"]), disabled=has_attempts)
        show_average = st.checkbox("Show class average", bool(quiz["show_average"]), disabled=has_attempts)
        randomize_questions = st.checkbox("Randomize question order", bool(quiz["randomize_questions"]), disabled=has_attempts)
        randomize_answers = st.checkbox("Randomize answer order", bool(quiz["randomize_answers"]), disabled=has_attempts)
        saved = st.form_submit_button("Save settings", type="primary", disabled=has_attempts, width="stretch")
    if saved:
        opening_value = datetime.combine(opening_day, opening_clock, tzinfo=timezone.utc)
        closing_value = datetime.combine(closing_day, closing_clock, tzinfo=timezone.utc)
        if opening_enabled and closing_enabled and closing_value <= opening_value:
            st.error("Closing time must be after opening time.")
        elif closing_enabled and closing_value <= datetime.now(timezone.utc):
            st.error("Closing time is in the past; students won't be able to take this quiz. Set a closing time in the future.")
        else:
            update_quiz_settings(quiz["id"], duration, passing, allow_retake, show_average, opening_value.isoformat(), closing_value.isoformat(), opening_enabled, closing_enabled, randomize_questions, randomize_answers)
            st.success("Settings saved.")


def assignment_editor(quiz) -> None:
    roster = students(quiz["owner_id"]); current = assigned_student_ids(quiz["id"])
    st.write("Choose who can see this assessment")
    st.caption("An empty selection means the quiz is available to every student. Selected students see it only on their dashboard.")
    audience_mode = st.radio("Assign to", ["Students", "Teams"], horizontal=True, key=f"assigned-mode-{quiz['id']}")
    if audience_mode == "Students":
        selected = st.multiselect("Assigned students", options=roster, default=[row for row in roster if row["id"] in current], format_func=lambda row: f"{row['name']}  ·  {row['email']}", key=f"assigned-{quiz['id']}")
        selected_ids = [row["id"] for row in selected]
    else:
        selected_teams = st.multiselect("Assigned teams", options=teams_for_teacher(quiz["owner_id"]), format_func=lambda team: team["name"], key=f"assigned-teams-{quiz['id']}")
        selected_ids = list(student_ids_for_teams(quiz["owner_id"], [team["id"] for team in selected_teams]))
    if st.button("Save assignment", type="primary", key=f"assign-save-{quiz['id']}"):
        set_quiz_assignments(quiz["id"], selected_ids); st.success("Assignment updated.")


def results(quiz) -> None:
    progress = student_progress_for_quiz(quiz["owner_id"], quiz["id"])
    if not progress:
        st.info("No students are assigned to this exam yet."); return
    frame = pd.DataFrame([{"Student": row["student"], "Email": row["email"], "Status": row["status"], "Score": row["score"], "Result": row["result"], "Last activity": row["last_activity"]} for row in progress])
    st.dataframe(frame, width="stretch", hide_index=True)
    st.download_button("Download student results CSV", frame.to_csv(index=False), "student-results.csv", "text/csv", key=f"results-{quiz['id']}")


def roster_page(user) -> None:
    roster = students(user["id"])
    performance = student_analytics(user["id"])
    teams = teams_for_teacher(user["id"])
    st.markdown('<div class="eyebrow">Teacher workspace</div><h1>Student roster</h1>', unsafe_allow_html=True)
    st.caption("Add the students you teach here. Only this roster appears in your assignment controls.")
    team_options = {team["id"]: team["name"] for team in teams}
    with st.container(border=True):
        st.subheader("Add an existing student to roster")
        lookup = st.text_input("Search student", placeholder="Search by name or email", key="existing-student-search")
        matches = [row for row in students() if not lookup.strip() or lookup.lower() in row["name"].lower() or lookup.lower() in row["email"].lower()]
        if matches:
            existing = st.selectbox("Find student", matches, format_func=lambda row: f"{row['name']} · {row['email']}", key="existing-student")
            existing_team_id = st.selectbox("Add to team", [None, *team_options], format_func=lambda value: "No team" if value is None else team_options[value], key="existing-student-team")
            if st.button("Add an existing student to roster", type="primary", width="stretch"):
                add_student_to_roster(user["id"], existing["name"], existing["email"], existing_team_id)
                st.success("Student added to your roster.")
                st.rerun()
        elif lookup.strip():
            st.info("No existing student matches that search.")
    with st.container(border=True):
        st.subheader("Add a new student")
        with st.form("add-student"):
            name = st.text_input("Student name", placeholder="e.g. Jordan Lee")
            email = st.text_input("Student email", placeholder="student@example.com")
            team_id = st.selectbox("Add to team", [None, *team_options], format_func=lambda value: "No team" if value is None else team_options[value], key="new-student-team")
            submitted = st.form_submit_button("Add a new student", type="primary", width="stretch")
        if submitted:
            if not name.strip() or "@" not in email:
                st.error("Enter a student name and a valid email address.")
            else:
                try:
                    add_student_to_roster(user["id"], name, email, team_id)
                except ValueError as exc:
                    st.error(str(exc))
                else:
                    st.success(f"{name.strip()} was added to your roster."); st.rerun()
    with st.container(border=True):
        st.subheader("Teams")
        st.caption("Choose a team to see its members. Use the optional search to add someone without opening the member list.")
        for team in teams:
            with st.expander(f"{team['name']} · {len(team_student_ids(team['id']))} members"):
                members = st.multiselect("Members", roster, default=[row for row in roster if row["id"] in team_student_ids(team["id"])], format_func=lambda row: f"{row['name']} · {row['email']}", key=f"team-members-{team['id']}")
                show_search = st.checkbox("Show search to add a member", key=f"show-team-search-{team['id']}")
                if show_search:
                    member_search = st.text_input("Search roster", placeholder="Search by name or email", key=f"team-search-{team['id']}")
                    matches = [row for row in roster if not member_search.strip() or member_search.lower() in row["name"].lower() or member_search.lower() in row["email"].lower()]
                    if matches:
                        candidate = st.radio("Add member", matches, format_func=lambda row: f"{row['name']} · {row['email']}", key=f"team-candidate-{team['id']}")
                        if st.button("Add member", key=f"add-team-member-{team['id']}"):
                            members = [*members, candidate] if candidate["id"] not in {row["id"] for row in members} else members
                            set_team_members(team["id"], [row["id"] for row in members])
                            st.rerun()
                if st.button("Save members", key=f"save-team-{team['id']}"):
                    set_team_members(team["id"], [row["id"] for row in members])
                    st.success("Team updated.")
    with st.container(border=True):
        st.subheader(f"Your Roster (of students) · {len(roster)}")
        if roster:
            stats = [{"Student": row["name"], "Email": row["email"], "Team": ", ".join(team["name"] for team in teams_for_student(user["id"], row["id"])), "Assigned": row["assigned_quizzes"], "Attempts": row["attempts"], "Completed": row["completed"], "Average": f"{row['average_score']:.1f}%" if row["average_score"] is not None else "-", "Pass rate": f"{row['pass_rate'] * 100:.0f}%" if row["pass_rate"] is not None else "-"} for row in performance]
            st.dataframe(pd.DataFrame(stats), width="stretch", hide_index=True)
            st.caption("Select a student below to view detailed analytics.")
            for row in performance:
                with st.container(border=True):
                    details, action = st.columns([5, 1])
                    details.write(f"**{row['name']}**  ·  {row['email']}")
                    details.caption(f"{row['completed']} completed · {row['average_score']:.1f}% average" if row["average_score"] is not None else "No completed tests yet")
                    if action.button("Details", key=f"details-{row['id']}", type="primary", width="stretch"):
                        st.session_state.detail_student_id = row["id"]
                        st.session_state.show_student_detail = True
            if st.session_state.pop("show_student_detail", False):
                student_detail_dialog(user, st.session_state.detail_student_id)
        else:
            st.info("Your roster is empty. Add a student above before assigning a quiz.")


@st.dialog("Student analytics")
def student_detail_dialog(user, student_id: int) -> None:
    detail = student_detail_analytics(user["id"], student_id)
    if not detail:
        st.error("Student is not in your roster.")
        return
    student = detail["student"]
    st.subheader(student["name"])
    st.caption(student["email"])
    st.write("Teams: " + (", ".join(team["name"] for team in detail["teams"]) or "No team"))
    results_data = [{"Test": row["title"], "Status": "Completed" if row["submitted_at"] else "In progress", "Score": f"{row['score_percent']:.1f}%" if row["score_percent"] is not None else "-", "Result": "Passed" if row["passed"] else ("Failed" if row["passed"] is not None else "-"), "Last activity": row["submitted_at"] or row["started_at"]} for row in detail["results"]]
    if results_data:
        st.dataframe(pd.DataFrame(results_data), width="stretch", hide_index=True)
    else:
        st.info("No test results yet.")
